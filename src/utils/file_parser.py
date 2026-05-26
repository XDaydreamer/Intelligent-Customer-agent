import os
import io
from pathlib import Path

SUPPORTED_EXTENSIONS = {".txt", ".json", ".docx", ".csv", ".xlsx", ".md"}


class FileParser:
    @staticmethod
    async def parse(filepath: str | Path) -> str:
        filepath = Path(filepath)
        ext = filepath.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        if ext == ".txt" or ext == ".md":
            return await FileParser._read_text(filepath)
        elif ext == ".json":
            import json
            content = await FileParser._read_text(filepath)
            data = json.loads(content)
            return json.dumps(data, ensure_ascii=False, indent=2)
        elif ext == ".docx":
            return await FileParser._read_docx(filepath)
        elif ext == ".csv":
            return await FileParser._read_csv(filepath)
        elif ext == ".xlsx":
            return await FileParser._read_xlsx(filepath)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    @staticmethod
    async def _read_text(filepath: Path) -> str:
        import aiofiles
        async with aiofiles.open(filepath, "r", encoding="utf-8") as f:
            return await f.read()

    @staticmethod
    async def _read_docx(filepath: Path) -> str:
        from docx import Document
        doc = Document(str(filepath))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)

    @staticmethod
    async def _read_csv(filepath: Path) -> str:
        import pandas as pd
        df = pd.read_csv(filepath)
        return df.to_string(index=False)

    @staticmethod
    async def _read_xlsx(filepath: Path) -> str:
        import pandas as pd
        df = pd.read_excel(filepath)
        return df.to_string(index=False)

    @staticmethod
    def parse_bytes(content: bytes, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        if ext == ".txt" or ext == ".md":
            return content.decode("utf-8")
        elif ext == ".json":
            import json
            data = json.loads(content.decode("utf-8"))
            return json.dumps(data, ensure_ascii=False, indent=2)
        elif ext == ".docx":
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n".join(paragraphs)
        elif ext == ".csv":
            import pandas as pd
            df = pd.read_csv(io.BytesIO(content))
            return df.to_string(index=False)
        elif ext == ".xlsx":
            import pandas as pd
            df = pd.read_excel(io.BytesIO(content))
            return df.to_string(index=False)
        raise ValueError(f"不支持的文件格式: {ext}")


async def parse_file(filepath: str | Path) -> str:
    return await FileParser.parse(filepath)
